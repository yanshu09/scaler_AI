import os
import docx
from redactor import PIIRedactor

def process_docx(input_path: str, output_path: str, redactor: PIIRedactor):
    """
    Reads a .docx file, extracts text from paragraphs and tables, 
    applies the PII redaction engine, and saves a new file 
    while preserving original formatting.
    """
    print(f"Loading document: {input_path}...")
    doc = docx.Document(input_path)
    
    print("Redacting normal paragraphs...")
    # Step 1: Process regular body paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            # Update paragraph text with redacted version
            para.text = redactor.redact_all(para.text)
            
    print("Redacting tables...")
    # Step 2: Traverse deeply into tables (Rows -> Cells -> Paragraphs)
    # This ensures PII inside structured data (like Prospectus tables) is caught
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text.strip():
                        para.text = redactor.redact_all(para.text)
                        
    print(f"Saving redacted document to: {output_path}...")
    doc.save(output_path)
    print("Done! ✅")

if __name__ == "__main__":
    # Initialize our core redactor engine
    redactor = PIIRedactor()
    
    # Define file paths
    input_file = "Red Herring Prospectus.docx"
    output_file = "Redacted_Red_Herring_Prospectus.docx"
    
    # Validation check to ensure file exists before processing
    if os.path.exists(input_file):
        process_docx(input_file, output_file, redactor)
    else:
        print(f"Error: '{input_file}' not found in the current directory.")